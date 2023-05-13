import io
import os
import docx
import openai
import PyPDF4

# Define o caminho do arquivo PDF
current_dir = os.getcwd()
pdf_path = os.path.join(current_dir, 'scriptspy' ,'peticaoraiz.pdf')

# Define as perguntas para a API da OpenAI
questions = [
    "Qual é o nome do autor da petição?",
    "Qual é o número do processo?",
    "Qual é a vara do processo?",
    "Qual é o objeto da ação?",
    "Qual é o pedido do autor?",
    "Qual é o fundamento jurídico do pedido?",
]

# Configuração da API da OpenAI
openai.api_key = "CODIGO API AQUI"

# Função para extrair o texto do PDF
def extract_text_from_pdf(pdf_path):
    with open(pdf_path, "rb") as f:
        pdf_reader = PyPDF4.PdfFileReader(f)
        page = pdf_reader.getPage(0)
        return page.extractText()

# Função para criar o documento Word e adicionar o texto extraído e as respostas
def create_word_document(text, answers):
    doc = docx.Document()
    doc.add_paragraph(text)
    for i in range(len(questions)):
        if i >= len(answers):
            doc.add_paragraph("Não foi possível obter resposta para a pergunta: " + questions[i])
        else:
            doc.add_paragraph(questions[i] + " " + answers[i])
    doc.save("relatorio-script8.docx")

# Extrai o texto do PDF
pdf_text = extract_text_from_pdf(pdf_path)

# Envia as perguntas para a API da OpenAI
openai_answers = openai.Completion.create(
    engine="davinci",
    prompt="\n".join(questions) + "\nTexto: " + pdf_text,
    temperature=0.7,
    max_tokens=200,
    n=1,
    stop=None,
    timeout=20,
)

# Armazena as respostas da API em uma lista
answers = []
for choice in openai_answers.choices:
    text = choice.text.strip()
    if text == "":
        answers.append("Não foi possível obter resposta para esta pergunta.")
    else:
        answers.append(text)

# Cria o documento Word com as informações extraídas e as respostas da API
create_word_document(pdf_text, answers)

print("Relatório criado com sucesso!")
