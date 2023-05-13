#Script Extrai do docx 
import io
import os
import docx
import openai

# Define o caminho do arquivo docx
current_dir = os.getcwd()
docx_path = os.path.join(current_dir, 'scriptspy' ,'peticaoraiz.docx')

# Define as perguntas para a API da OpenAI
questions = [
    "Qual é o nome completo do autor da petição, incluindo seu número de CPF e endereço atual?",
    "Qual é o número completo do processo, incluindo o ano, a comarca e a vara responsável pelo julgamento?",
    "Qual é o número de protocolo da petição e a data em que foi protocolada?",
    "Qual é o objeto específico da ação, descrevendo em detalhes as circunstâncias que levaram à abertura do processo?",
    "Qual é o pedido do autor em relação à ação, incluindo os valores monetários solicitados e os prazos para cumprimento?",
    "Qual é o fundamento jurídico do pedido, incluindo a legislação aplicável e as jurisprudências que sustentam a ação?"
]

# Configuração da API da OpenAI
openai.api_key = "CODIGO API AQUI"

# Função para extrair o texto do arquivo docx
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Função para criar o documento Word e adicionar o texto extraído e as respostas
def create_word_document(text, answers):
    doc = docx.Document()
    doc.add_paragraph(text)
    for i in range(len(questions)):
        if i >= len(answers):
            doc.add_paragraph("Não foi possível obter resposta para a pergunta: " + questions[i])
        else:
            doc.add_paragraph(questions[i] + " " + answers[i])
    doc.save("relatorio-script9.docx")

# Extrai o texto do arquivo docx
docx_text = extract_text_from_docx(docx_path)

# Envia as perguntas para a API da OpenAI
openai_answers = openai.Completion.create(
    engine="davinci",
    prompt="\n".join(questions) + "\nTexto: " + docx_text,
    temperature=0.7,
    max_tokens=200,
    n=6,
    stop=None,
    timeout=20,
)

# Armazena as respostas da API em uma lista
answers = []
for choice in openai_answers.choices:
    text = choice.text.strip()
    answers.append(text)

# Cria o documento Word com as informações extraídas e as respostas da API
create_word_document(docx_text, answers)

print("Relatório criado com sucesso!")
