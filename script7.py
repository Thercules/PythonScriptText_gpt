import requests
import os
from docx import Document
import openai
import PyPDF4
import io

openai.api_key = "CODIGO API AQUI"
model_engine = "text-davinci-002"

# obtém o diretório atual do script
current_dir = os.path.dirname(os.path.abspath(__file__))

# monta o caminho para o arquivo PDF
pdf_path = os.path.join(current_dir, 'dados', 'peticao.pdf')

# monta o caminho para o arquivo DOCX
docx_path = os.path.join(current_dir, 'dados', 'peticao.docx')

# Extrai o texto do arquivo DOCX usando a biblioteca python-docx
try:
    doc = Document(docx_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
except FileNotFoundError as err:
    print(f"Arquivo não encontrado: {err}")
    exit(1)

try:
    # Realiza a requisição para a API da OpenAI
    response = requests.post("https://api.openai.com/v1/engines/" + model_engine + "/completions", headers={
        "Content-Type": "application/json",
        "Authorization": f"Bearer {openai.api_key}"
    }, json={
        "prompt": text,
        "temperature": 0.7,
        "max_tokens": 3500,
        "top_p": 1,
        "frequency_penalty": 0,
        "presence_penalty": 0
    })

    # Obtém a interpretação da petição gerada pela OpenAI
    interpretation = response.json()["choices"][0]["text"]
    
    # Salva a interpretação em um arquivo Word
    document = Document()
    document.add_paragraph(f"Interpretação da Petição: {interpretation}")
    document.save("relatorio.docx")

except requests.exceptions.HTTPError as err:
    print(f"Erro HTTP: {err}")
    print(f"Resposta da API: {err.response.content}")
except KeyError as err:
    print(f"Erro ao obter interpretação: {err}")
    print(f"Resposta da API: {response.content}")
except Exception as err:
    print(f"Erro não esperado: {err}")
